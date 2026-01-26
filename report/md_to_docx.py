
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default style to Arial 12 (approximation)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block: # End of block
                continue
            else: # Start of block
                continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(10)
            continue

        # Handle Tables (Basic)
        if '|' in line and (line.startswith('|') or line.endswith('|')):
            table_lines.append(line)
            continue
        else:
            if table_lines:
                # Process collected table
                process_table(doc, table_lines)
                table_lines = []

        if not line:
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line[0].isdigit() and line[1] == '.':
            p = doc.add_paragraph(line, style='List Number')
        # Images/Captions (Markdown syntax ![alt](src))
        elif line.startswith('![') and '](' in line:
            # Just put the alt text as a placeholder caption
            alt_text = line.split('![')[1].split('](')[0]
            p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {alt_text}]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.italic = True
        # Standard Paragraph
        else:
            # Handle Bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Flush any remaining table
    if table_lines:
        process_table(doc, table_lines)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def process_table(doc, lines):
    # Determine rows and cols
    # Filter out separator lines like |---|---|
    data_rows = [l for l in lines if not set(l.strip('| ')).issubset({'-', ':'})]
    
    if not data_rows:
        return

    # Count columns based on first row
    first_row_cells = [c.strip() for c in data_rows[0].strip('|').split('|')]
    cols = len(first_row_cells)
    
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    
    for row_line in data_rows:
        cells_data = [c.strip() for c in row_line.strip('|').split('|')]
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(cells_data):
            if i < len(row_cells):
                row_cells[i].text = cell_text

if __name__ == '__main__':
    parse_markdown_to_docx('AFASA_2.0_Final_Report.md', 'AFASA_2.0_Final_Report.docx')
